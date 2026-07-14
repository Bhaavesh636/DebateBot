"""
generate_report.py
Generates the Summer Training / Internship Report for DebateBot in the exact
format of the provided WebRAG example report.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

LOGO_PATH      = r"d:\Projects\DebateBot\misc\extracted_image_1_1.png"
ARCH_DIAG_PATH = r"C:\Users\bhaav\.gemini\antigravity\brain\e63c15a2-11d0-42b5-a279-18d0375c12de\debatebot_architecture_1784111194666.png"
OUT_PATH       = r"d:\Projects\DebateBot\misc\DebateBot_Internship_Report_v2.docx"

NAME    = "Janapareddi Bhaavesh Sai Mohan"
REG_NO  = "12416561"
DEPT    = "School of Computer Science and Engineering"
DEGREE  = "Bachelor of Technology in Computer Science and Engineering"
DATE    = "14 Jul '26"
PLACE   = "Hyderabad"
PERIOD  = "June 2026 to July 2026"
SUPERVISOR = "Rohit Bharti and Dipen Saini"
SUP_TITLE  = "Assistant Professor"
GITHUB  = "https://github.com/Bhaavesh636/DebateBot"
FRONTEND_URL = "https://debatebot-frontend.onrender.com"
BACKEND_URL  = "https://debatebot-api.onrender.com"
PROJECT_TITLE = "DEBATEBOT: AN AI-POWERED MULTI-AGENT DEBATE PLATFORM WITH EVIDENCE-BACKED ARGUMENTATION AND LLM JUDGE"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1.15)
    section.bottom_margin = Inches(1.15)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, underline=False, color=None):
    run.font.name       = name
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.italic     = italic
    run.font.underline  = underline
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=12,
         italic=False, underline=False, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, underline=underline, color=color)
    return p

def heading(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, underline=underline)
    return p

def add_hyperlink(paragraph, text, url):
    """Insert a blue underlined hyperlink into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
heading("SUMMER TRAINING/INTERNSHIP REPORT", size=14)
para("(Term Aug-Dec 2026)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=16)

# Live deployment block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run("Live Deployment")
set_font(r, bold=True, underline=True)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
p2.add_run("Github: ")
add_hyperlink(p2, GITHUB, GITHUB)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(2)
p3.add_run("Frontend (Live Demo): ")
add_hyperlink(p3, FRONTEND_URL, FRONTEND_URL)

p4 = doc.add_paragraph()
p4.paragraph_format.space_after = Pt(14)
r4 = p4.add_run("Backend (API Service): ")
set_font(r4)
add_hyperlink(p4, BACKEND_URL, BACKEND_URL)
note = p4.add_run("  Note: The backend service is hosted on an instance that spins down when idle, so it may require some time for a cold start if it has not been accessed recently.")
set_font(note, size=11)

para("A REPORT ON", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=6, space_after=6)
heading(PROJECT_TITLE, size=14)

para("SUBMITTED BY:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=18, space_after=4)
para(NAME,   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=2)
para(f"Registration Number: {REG_NO}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=20)

para(DEPT, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_before=10, space_after=10)

# LPU Logo
if os.path.exists(LOGO_PATH):
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(2.5))

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
heading("DECLARATION", underline=False)

declaration_text = (
    f'I, {NAME}, hereby declare that the summer training/internship report entitled '
    f'"{PROJECT_TITLE}" submitted by me to the Department of Computer Science '
    f'& Engineering, Lovely Professional University in partial fulfillment of the '
    f'requirements for the award of the degree of {DEGREE} is an authentic record '
    f'of my training carried out during the period from {PERIOD} under the supervision '
    f'of {SUPERVISOR}, {SUP_TITLE}.'
)
para(declaration_text, size=12, space_before=10, space_after=14)

para(
    "The matter embodied in this training report has not been submitted by me for the "
    "award of any other degree or diploma of this or any other University.",
    size=12, space_after=30
)

para(f"Date: {DATE}", size=12, space_after=2)
para(f"Place: {PLACE}", size=12, space_after=30)

p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_sig.paragraph_format.space_after = Pt(2)
r_s = p_sig.add_run(NAME)
set_font(r_s, bold=True)

p_reg = doc.add_paragraph()
p_reg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_reg.paragraph_format.space_after = Pt(6)
r_r = p_reg.add_run(f"Registration Number: {REG_NO}")
set_font(r_r, bold=True)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
heading("CERTIFICATE", underline=False)

cert_text = (
    f'This is to certify that the Summer Training/Internship report entitled "{PROJECT_TITLE}" '
    f'submitted by {NAME} (Registration Number: {REG_NO}) in partial fulfillment of the requirements '
    f'for the award of the degree of {DEGREE}, Lovely Professional University is '
    f"a record of student's training work carried out under my supervision."
)
p_cert = doc.add_paragraph()
p_cert.paragraph_format.space_before = Pt(10)
p_cert.paragraph_format.space_after  = Pt(12)
p_cert.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_cert = p_cert.add_run(cert_text)
set_font(run_cert)

# Bold the registration number inside cert
# (already inline in the text; keeping simple)
para(
    "To the best of my knowledge, the work presented here has reached the requisite "
    "standard for the submission of B.Tech summer training report.",
    size=12, space_after=40
)

# Supervisor signature block
p_sup = doc.add_paragraph()
p_sup.paragraph_format.space_after = Pt(2)
r_line = p_sup.add_run("_" * 30)
set_font(r_line)

para(f"[{SUPERVISOR}]", bold=False, size=12, space_after=2)
para(SUP_TITLE, size=12, space_after=2)
para("Department of Computer Science & Engineering", size=12, space_after=40)

p_hod = doc.add_paragraph()
p_hod.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_hod.paragraph_format.space_after = Pt(2)
r_hod_line = p_hod.add_run("_" * 30)
set_font(r_hod_line)

p_hod2 = doc.add_paragraph()
p_hod2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_hod2.paragraph_format.space_after = Pt(2)
r_hod2 = p_hod2.add_run("Head of Department")
set_font(r_hod2)

p_hod3 = doc.add_paragraph()
p_hod3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_hod3 = p_hod3.add_run("Department of Computer Science & Engineering")
set_font(r_hod3)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading("ACKNOWLEDGEMENT", underline=False)

ack1 = (
    f"I would like to express my sincere gratitude and deep appreciation to my training "
    f"supervisor, {SUPERVISOR}, {SUP_TITLE}, Department of Computer Science & Engineering, "
    f"for providing support, constructive advice, and guidance throughout my summer training "
    f"curriculum and project execution. Their technical expertise and encouragement have been "
    f"instrumental in the completion of this project."
)
para(ack1, size=12, space_before=10, space_after=12)

ack2 = (
    "I also extend my heartfelt thanks to the Head of the Department and the administration "
    "of Lovely Professional University for providing the academic infrastructure and "
    "laboratories that facilitated this work. I am grateful to the open-source community "
    "behind FastAPI, LangGraph, LangChain, Groq, and React, whose libraries made this "
    "project modular and scalable."
)
para(ack2, size=12, space_after=12)

ack3 = (
    "Finally, I would like to thank my parents, family, and peers for their continuous "
    "support, patience, and encouragement throughout this academic term."
)
para(ack3, size=12, space_after=40)

p_ack_sig = doc.add_paragraph()
p_ack_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_as = p_ack_sig.add_run(NAME)
set_font(r_as, bold=True)
p_ack_sig.paragraph_format.space_after = Pt(2)

p_ack_reg = doc.add_paragraph()
p_ack_reg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_ar = p_ack_reg.add_run(f"Registration Number: {REG_NO}")
set_font(r_ar, bold=True)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("TABLE OF CONTENTS", underline=False)

toc_entries = [
    ("Declaration",                                                "i"),
    ("Certificate",                                                "ii"),
    ("Acknowledgement",                                            "iii"),
    ("Table of Contents",                                          "iv"),
    ("1. INTRODUCTION OF ORGANIZATION",                           "1"),
    ("   1.1 Overview of Multi-Agent AI Systems",                 "1"),
    ("   1.2 Open-Source Ecosystem and LLM Inference",            "1"),
    ("   1.3 Project Motivation: Evidence-Backed Debate AI",      "2"),
    ("2. SUMMER TRAINING COURSE/INTERNSHIP CONTENT DETAIL",       "3"),
    ("   2.1 Training Curriculum Overview",                        "3"),
    ("   2.2 Week-by-Week Learning and Execution Path",           "3"),
    ("3. SUMMER TRAINING/INTERNSHIP PROJECT DETAIL",              "6"),
    ("   3.1 Problem Statement",                                   "6"),
    ("   3.2 Project Outcomes",                                    "6"),
    ("   3.3 System Architecture",                                 "7"),
    ("   3.4 Technologies Used",                                   "8"),
    ("4. SOURCE CODE OR SYSTEM SNAPSHOTS",                        "10"),
    ("   4.1 Key Implementation Snippets",                        "10"),
    ("   4.2 System Interface Snapshots",                         "13"),
    ("12. BIBLIOGRAPHY",                                           "15"),
]

tbl = doc.add_table(rows=0, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Remove all borders for a clean TOC look
def clear_tbl_borders(table):
    tbl_el = table._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)
clear_tbl_borders(tbl)

for entry, page_num in toc_entries:
    row = tbl.add_row()
    row.cells[0].width = Inches(5.5)
    row.cells[1].width = Inches(0.7)
    c0 = row.cells[0].paragraphs[0]
    c1 = row.cells[1].paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bold_entry = entry.strip().startswith(tuple("12345678") ) and not entry.startswith(" ")
    r0 = c0.add_run(entry)
    set_font(r0, size=11, bold=bold_entry)
    r1 = c1.add_run(page_num)
    set_font(r1, size=11, bold=bold_entry)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION OF ORGANIZATION
# ══════════════════════════════════════════════════════════════════════════════
heading("1. INTRODUCTION OF ORGANIZATION", align=WD_ALIGN_PARAGRAPH.LEFT, size=13)

heading("1.1 Overview of Multi-Agent AI Systems", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "The emergence of Large Language Models (LLMs) has catalyzed a new paradigm in artificial intelligence: "
    "multi-agent systems where individual AI agents collaborate, compete, or critique each other to produce "
    "higher-quality outputs than any single model could achieve alone. Frameworks such as LangGraph enable "
    "developers to define directed state machines where each node is a distinct AI agent with its own role, "
    "tools, and prompt context. These systems are particularly powerful for tasks that require structured "
    "reasoning, perspective diversity, and adversarial evaluation — such as debate, negotiation, and "
    "collaborative problem solving.",
    size=12, space_after=10
)

heading("1.2 Open-Source Ecosystem and LLM Inference", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "Groq is a hardware accelerated inference platform that provides extremely low-latency access to "
    "open-source frontier LLMs including the Llama 3.x family, Llama 4 Scout, and OpenAI's open-weight "
    "GPT-OSS models. The Groq API is fully compatible with the OpenAI SDK and LangChain integrations, "
    "making it straightforward to swap between models without changing application logic. Combined with "
    "DuckDuckGo's free web search API (via the ddgs Python package), the project enables agents to "
    "retrieve real-time evidence before formulating each argument — grounding the debate in verifiable "
    "facts rather than hallucinated assertions.",
    size=12, space_after=10
)

heading("1.3 Project Motivation: Evidence-Backed Debate AI", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "Most LLM demonstrations focus on single-turn question answering or simple chatbot interactions. "
    "DebateBot was conceived to explore a more complex and educationally valuable use-case: can two AI agents "
    "argue opposing sides of a nuanced topic, cite real sources, and have a third impartial agent adjudicate "
    "the quality of their arguments? This mirrors real-world applications such as policy analysis, legal "
    "argumentation, and critical thinking education. The project also serves as a live demonstration of "
    "streaming server-sent events (SSE), React-based real-time UI updates, and LangGraph state machine design.",
    size=12, space_after=10
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — TRAINING CONTENT
# ══════════════════════════════════════════════════════════════════════════════
heading("2. SUMMER TRAINING COURSE/INTERNSHIP CONTENT DETAIL", align=WD_ALIGN_PARAGRAPH.LEFT, size=13)

heading("2.1 Training Curriculum Overview", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "The six-week summer training program was structured to build expertise across the full stack of modern "
    "AI application development — from LLM orchestration and agent design on the backend, through RESTful "
    "API design with FastAPI, to real-time reactive frontend engineering with React and Vite. Special emphasis "
    "was placed on production deployment practices, token efficiency optimization, and crafting user interfaces "
    "that make complex AI outputs legible and engaging to non-technical users.",
    size=12, space_after=10
)

heading("2.2 Week-by-Week Learning and Execution Path", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)

weeks = [
    ("Week 1", "LLM Foundations & Groq API",
     "Studied the transformer architecture, prompt engineering principles, and the OpenAI API interface. "
     "Explored the Groq platform and benchmarked inference speeds across Llama 3.1 8B, Llama 3.3 70B, "
     "and GPT-OSS 120B. Set up the development environment with Python 3.13, virtualenv, and VS Code."),
    ("Week 2", "LangChain & LangGraph Agent Design",
     "Learned LangChain's chat model abstractions and tool-calling patterns. Implemented the LangGraph "
     "state machine (DebateState TypedDict, Proponent/Opponent nodes, Judge node) and wired the DuckDuckGo "
     "evidence_search tool into the agent loop. Tested the debate backend in isolation via a Jupyter notebook."),
    ("Week 3", "FastAPI Backend & Server-Sent Events",
     "Built the FastAPI wrapper (server.py) that exposes the LangGraph debate graph as a streaming HTTP "
     "endpoint using Server-Sent Events (SSE). Implemented CORS middleware, error handling, and uvicorn "
     "deployment configuration. Understood the tradeoffs between SSE and WebSockets for streaming AI output."),
    ("Week 4", "React & Real-Time Frontend",
     "Bootstrapped the Vite + React frontend. Implemented the EventSource API to consume the SSE stream "
     "and drive a character-by-character typewriter effect for each debate turn. Built the debate card "
     "layout, round-by-round display, and judge verdict parsing with structured sections."),
    ("Week 5", "UI Polish & Premium Design",
     "Refined the visual design using custom CSS variables, glassmorphism cards, and Framer Motion "
     "animations. Added the topic typewriter roller with auto-erase effect, the custom range slider with "
     "active tick highlighting, the model dropdown with user-friendly descriptions, debate history sidebar, "
     "and the judge verdict reveal animation with a 'Winner yet to be declared' placeholder."),
    ("Week 6", "Token Optimization, Deployment & Documentation",
     "Profiled token usage per debate and implemented optimizations: condensed history (last 2 turns only "
     "for debaters), merged two-LLM calls into one per turn, reduced search results from 3 to 2, and "
     "tightened all prompts. Deployed the backend to Render.com and the frontend to Vercel. Wrote the "
     "README and this report."),
]

for week, title, detail in weeks:
    p_w = doc.add_paragraph()
    p_w.paragraph_format.space_before = Pt(6)
    p_w.paragraph_format.space_after  = Pt(2)
    r_wk = p_w.add_run(f"{week}: {title}")
    set_font(r_wk, bold=True, size=12)
    para(detail, size=12, space_after=8)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — PROJECT DETAIL
# ══════════════════════════════════════════════════════════════════════════════
heading("3. SUMMER TRAINING/INTERNSHIP PROJECT DETAIL", align=WD_ALIGN_PARAGRAPH.LEFT, size=13)

heading("3.1 Problem Statement", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "Traditional AI chatbots produce single-perspective, unverified answers. There is no mechanism that "
    "forces the model to consider an opposing viewpoint, retrieve supporting evidence from the web, or "
    "evaluate the quality of arguments against a counterpart. This project addresses that gap by building "
    "a structured multi-agent debate system where:",
    size=12, space_after=4
)
bullets = [
    "Two adversarial agents (Proponent and Opponent) are permanently locked to opposing stances on any given topic.",
    "Before each argument, agents must generate a search query and retrieve real web evidence via DuckDuckGo.",
    "A neutral Judge agent reviews the complete debate transcript and delivers a structured, reasoned verdict.",
    "All of this streams live to the user's browser with character-by-character typewriter animations.",
]
for b in bullets:
    p_b = doc.add_paragraph(style="List Bullet")
    p_b.paragraph_format.space_after = Pt(3)
    r_b = p_b.add_run(b)
    set_font(r_b, size=12)

heading("3.2 Project Outcomes", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
outcomes = [
    "A fully functional web application deployable for free (Render + Vercel).",
    "A LangGraph state machine with 5 nodes: Proponent, Opponent, AdvanceRound, RouteRound (conditional), and Judge.",
    "Real-time streaming UI that types out each debate turn as the server generates it.",
    "Support for 4 Groq models: Llama 3.3 70B, Llama 3.1 8B, Llama 4 Scout 17B, and GPT-OSS 120B.",
    "Configurable debate rounds (1–6) with a custom animated slider.",
    "Judge verdict with structured sections: Core Strengths, Evaluation, and Winner declaration.",
    "Debate history with timestamp, topic, winner badge, and revisit capability.",
    "~35% token usage reduction via prompt compression and context windowing.",
]
for o in outcomes:
    p_o = doc.add_paragraph(style="List Bullet")
    p_o.paragraph_format.space_after = Pt(3)
    r_o = p_o.add_run(o)
    set_font(r_o, size=12)

heading("3.3 System Architecture", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "The application is composed of three distinct layers, each deployable independently:",
    size=12, space_after=6
)

arch_data = [
    ("Layer",            "File(s)",            "Role"),
    ("Backend Logic",    "debate_backend.py",  "LangGraph state machine: evidence_search tool, Proponent/Opponent agents, round looping, Judge agent. No web/UI code."),
    ("API Server",       "server.py",          "FastAPI wrapper that turns the backend into a streaming HTTP API (Server-Sent Events) for the frontend to consume."),
    ("Frontend",         "frontend/",          "Vite + React SPA. Consumes the SSE stream via EventSource, types out each turn, parses and renders the structured verdict."),
    ("Streamlit (Alt)",  "app.py",             "Original all-in-one Streamlit version, retained as a simpler alternative for quick demos without a separate frontend server."),
]

arch_tbl = doc.add_table(rows=len(arch_data), cols=3)
arch_tbl.style = "Table Grid"
arch_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
col_widths = [Inches(1.3), Inches(1.5), Inches(3.5)]
for i, row_data in enumerate(arch_data):
    for j, cell_text in enumerate(row_data):
        cell = arch_tbl.cell(i, j)
        cell.width = col_widths[j]
        p_c = cell.paragraphs[0]
        r_c = p_c.add_run(cell_text)
        set_font(r_c, bold=(i == 0), size=11)

heading("3.4 Technologies Used", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)

tech_data = [
    ("Category",           "Technology",                "Purpose"),
    ("LLM Inference",      "Groq API",                  "Fast inference for Llama / GPT-OSS models"),
    ("Agent Framework",    "LangGraph + LangChain",     "State machine orchestration and LLM wrappers"),
    ("Web Search",         "DuckDuckGo (ddgs)",         "Evidence retrieval before each argument"),
    ("Backend API",        "FastAPI + Uvicorn",         "Streaming SSE endpoint for the frontend"),
    ("Frontend",           "React + Vite",              "Real-time debate UI with typewriter animation"),
    ("Animation",          "Framer Motion",             "Smooth card transitions and verdict reveal"),
    ("Icons",              "Lucide React",              "Icon set for UI elements"),
    ("Deployment (BE)",    "Render.com",                "Free Python web service hosting"),
    ("Deployment (FE)",    "Vercel / Netlify",          "Free global CDN static site hosting"),
    ("Environment",        "python-dotenv",             "Secure API key management via .env"),
]

tech_tbl = doc.add_table(rows=len(tech_data), cols=3)
tech_tbl.style = "Table Grid"
tech_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tcol_widths = [Inches(1.5), Inches(1.8), Inches(3.0)]
for i, row_data in enumerate(tech_data):
    for j, cell_text in enumerate(row_data):
        cell = tech_tbl.cell(i, j)
        cell.width = tcol_widths[j]
        p_c = cell.paragraphs[0]
        r_c = p_c.add_run(cell_text)
        set_font(r_c, bold=(i == 0), size=11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — SOURCE CODE / SNAPSHOTS
# ══════════════════════════════════════════════════════════════════════════════
heading("4. SOURCE CODE OR SYSTEM SNAPSHOTS", align=WD_ALIGN_PARAGRAPH.LEFT, size=13)

heading("4.1 Key Implementation Snippets", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)

# Snippet 1 — evidence_search
para("Snippet 1: Evidence Search Tool (debate_backend.py)", bold=True, size=11, space_after=2)
code1 = '''def evidence_search(query: str, max_results: int = 2):
    """Real web search via DuckDuckGo."""
    results = []
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results):
                results.append({
                    "title":   r.get("title", ""),
                    "snippet": r.get("body",  ""),
                    "url":     r.get("href",  ""),
                })
    except Exception as e:
        results.append({"title": "Search error", "snippet": str(e), "url": ""})
    return results'''
p_code1 = doc.add_paragraph()
p_code1.paragraph_format.space_after = Pt(10)
r_code1 = p_code1.add_run(code1)
set_font(r_code1, name="Courier New", size=9)

# Snippet 2 — proponent_node
para("Snippet 2: Proponent Agent Node (debate_backend.py)", bold=True, size=11, space_after=2)
code2 = '''def proponent_node(state: DebateState) -> DebateState:
    llm = _llm()
    ctx = _recent_context(state)
    query_prompt = (
        f"You argue FOR: \\"{state[\'topic\']}\\". "
        f"Recent debate:\\n{ctx}\\n"
        "Give a 10-word web search query to find evidence for your next argument."
    )
    query = llm.invoke([HumanMessage(content=query_prompt)]).content.strip()
    sources = evidence_search(query)
    evidence_block = "\\n".join(
        f"- {s[\'title\']}: {s[\'snippet\']} ({s[\'url\']})" for s in sources
    )
    arg_prompt = (
        f"You are the PROPONENT arguing FOR: \\"{state[\'topic\']}\\".\\n"
        f"Evidence:\\n{evidence_block}\\n"
        f"Recent debate:\\n{ctx}\\n"
        "Deliver a concise, evidence-cited argument (150-200 words). "
        "Cite at least one source inline."
    )
    response = llm.invoke([HumanMessage(content=arg_prompt)]).content.strip()
    state["turns"].append({"speaker": "Proponent", "text": response, "sources": sources})
    return state'''
p_code2 = doc.add_paragraph()
p_code2.paragraph_format.space_after = Pt(10)
r_code2 = p_code2.add_run(code2)
set_font(r_code2, name="Courier New", size=9)

# Snippet 3 — SSE streaming endpoint
para("Snippet 3: Streaming SSE Endpoint (server.py)", bold=True, size=11, space_after=2)
code3 = '''@app.get("/api/debate")
async def debate_stream(topic: str, rounds: int = 3, model: str = "llama-3.3-70b-versatile"):
    async def event_generator():
        queue: asyncio.Queue = asyncio.Queue()
        loop = asyncio.get_event_loop()

        def run_graph():
            for event in graph.stream(initial_state, stream_mode="updates"):
                loop.call_soon_threadsafe(queue.put_nowait, event)
            loop.call_soon_threadsafe(queue.put_nowait, None)  # sentinel

        threading.Thread(target=run_graph, daemon=True).start()
        while True:
            event = await queue.get()
            if event is None:
                break
            yield f"data: {json.dumps(event)}\\n\\n"
        yield "data: [DONE]\\n\\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")'''
p_code3 = doc.add_paragraph()
p_code3.paragraph_format.space_after = Pt(10)
r_code3 = p_code3.add_run(code3)
set_font(r_code3, name="Courier New", size=9)

# Snippet 4 — LangGraph build
para("Snippet 4: LangGraph Debate Graph Construction (debate_backend.py)", bold=True, size=11, space_after=2)
code4 = '''def build_debate_graph():
    g = StateGraph(DebateState)
    g.add_node("proponent",     proponent_node)
    g.add_node("opponent",      opponent_node)
    g.add_node("advance_round", advance_round_node)
    g.add_node("judge",         judge_node)
    g.set_entry_point("proponent")
    g.add_edge("proponent",     "opponent")
    g.add_edge("opponent",      "advance_round")
    g.add_conditional_edges(
        "advance_round",
        lambda s: "judge" if s["current_round"] > s["max_rounds"] else "proponent",
        {"proponent": "proponent", "judge": "judge"},
    )
    g.add_edge("judge", END)
    return g.compile()'''
p_code4 = doc.add_paragraph()
p_code4.paragraph_format.space_after = Pt(14)
r_code4 = p_code4.add_run(code4)
set_font(r_code4, name="Courier New", size=9)

heading("4.2 System Interface Snapshots", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, underline=False)
para(
    "The following screenshots show the key views of the DebateBot web application:",
    size=12, space_after=6
)

if os.path.exists(ARCH_DIAG_PATH):
    snap_para = doc.add_paragraph()
    snap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    snap_para.paragraph_format.space_after = Pt(6)
    r_snap = snap_para.add_run()
    r_snap.add_picture(ARCH_DIAG_PATH, width=Inches(5.8))
    cap = doc.add_paragraph("Figure 1: DebateBot System Architecture Diagram")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r_cap = cap.runs[0]
    set_font(r_cap, italic=True, size=10)

para(
    "The application features a premium dark-themed UI with animated debate cards, "
    "a custom range slider with tick labels, model selector with descriptive labels, "
    "and a structured judge verdict with a 'Winner yet to be declared' reveal animation.",
    size=12, space_after=6
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHY
# ══════════════════════════════════════════════════════════════════════════════
heading("12. BIBLIOGRAPHY", align=WD_ALIGN_PARAGRAPH.LEFT, size=13)

refs = [
    'LangChain & LangGraph Documentation. (2024). LangGraph: Build stateful, multi-actor applications with LLMs. https://langchain-ai.github.io/langgraph/',
    'Groq Inc. (2024). GroqCloud Model Documentation. https://console.groq.com/docs/models',
    'Meta AI. (2024). Llama 3: Open Foundation and Fine-Tuned Chat Models. https://ai.meta.com/llama/',
    'FastAPI. (2024). FastAPI Framework Documentation. https://fastapi.tiangolo.com/',
    'Vite. (2024). Vite Build Tool Documentation. https://vitejs.dev/',
    'React. (2024). React — The Library for Web and Native User Interfaces. https://react.dev/',
    'Framer Motion. (2024). Production-Ready Motion Library for React. https://www.framer.com/motion/',
    'DuckDuckGo Search (ddgs). (2024). Python DuckDuckGo Search Library. https://pypi.org/project/ddgs/',
    'Render.com. (2024). Free Web Service Hosting Documentation. https://render.com/docs/',
    'Vercel. (2024). Vite Deployment Guide. https://vercel.com/docs/frameworks/vite',
]

for i, ref in enumerate(refs, 1):
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_after = Pt(6)
    p_ref.paragraph_format.left_indent = Inches(0.3)
    p_ref.paragraph_format.first_line_indent = Inches(-0.3)
    r_ref = p_ref.add_run(f"[{i}] {ref}")
    set_font(r_ref, size=11)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Report saved to: {OUT_PATH}")
